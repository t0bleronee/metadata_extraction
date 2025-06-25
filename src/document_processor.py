"""
Document Processing Module for Automated Metadata Generation
This module handles text extraction from various document formats (PDF, DOCX, TXT)
with OCR support for scanned documents.
"""

import os
import io
import logging
from typing import Dict, List, Optional, Tuple
from pathlib import Path

# Document processing libraries
import PyPDF2
import pdfplumber
from docx import Document
import pytesseract
from PIL import Image
import fitz  # PyMuPDF for better PDF handling

# Text processing
import re
from textblob import TextBlob

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DocumentProcessor:
    """
    A comprehensive document processor that extracts text from various formats
    and prepares it for metadata generation.
    """
    
    def __init__(self, tesseract_path: Optional[str] = None):
        """
        Initialize the document processor.
        
        Args:
            tesseract_path: Path to tesseract executable (if not in PATH)
        """
        self.supported_formats = ['.pdf', '.docx', '.txt', '.doc']
        
        # Set tesseract path if provided
        if tesseract_path:
            pytesseract.pytesseract.tesseract_cmd = tesseract_path
        
        # Test OCR availability
        self._test_ocr()
    
    def _test_ocr(self) -> bool:
        """Test if OCR is working properly."""
        try:
            pytesseract.get_tesseract_version()
            logger.info("OCR (Tesseract) is available")
            return True
        except Exception as e:
            logger.warning(f"OCR not available: {e}")
            return False
    
    def process_document(self, file_path: str) -> Dict:
        """
        Process a document and extract text content with metadata.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Dictionary containing extracted text and basic metadata
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_path = Path(file_path)
        file_extension = file_path.suffix.lower()
        
        if file_extension not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_extension}")
        
        # Basic file metadata
        file_stats = file_path.stat()
        basic_metadata = {
            'filename': file_path.name,
            'file_size': file_stats.st_size,
            'file_extension': file_extension,
            'file_path': str(file_path)
        }
        
        # Extract text based on file type
        try:
            if file_extension == '.pdf':
                text_content = self._process_pdf(file_path)
            elif file_extension in ['.docx', '.doc']:
                text_content = self._process_docx(file_path)
            elif file_extension == '.txt':
                text_content = self._process_txt(file_path)
            else:
                raise ValueError(f"Handler not implemented for {file_extension}")
            
            # Process and clean text
            processed_text = self._clean_text(text_content)
            
            # Basic text statistics
            text_stats = self._get_text_statistics(processed_text)
            
            return {
                'raw_text': text_content,
                'processed_text': processed_text,
                'file_metadata': basic_metadata,
                'text_statistics': text_stats,
                'processing_status': 'success'
            }
            
        except Exception as e:
            logger.error(f"Error processing {file_path}: {e}")
            return {
                'raw_text': '',
                'processed_text': '',
                'file_metadata': basic_metadata,
                'text_statistics': {},
                'processing_status': 'failed',
                'error': str(e)
            }
    
    def _process_pdf(self, file_path: Path) -> str:
        """Extract text from PDF files with OCR fallback."""
        text_content = ""
        
        try:
            # Try pdfplumber first (better for structured PDFs)
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_content += page_text + "\n"
            
            # If we got very little text, try OCR
            if len(text_content.strip()) < 100:
                logger.info(f"Low text extraction from {file_path.name}, trying OCR...")
                text_content = self._process_pdf_with_ocr(file_path)
                
        except Exception as e:
            logger.warning(f"pdfplumber failed for {file_path.name}: {e}")
            # Fallback to PyPDF2
            try:
                text_content = self._process_pdf_pypdf2(file_path)
            except Exception as e2:
                logger.warning(f"PyPDF2 also failed: {e2}")
                # Last resort: OCR
                text_content = self._process_pdf_with_ocr(file_path)
        
        return text_content
    
    def _process_pdf_pypdf2(self, file_path: Path) -> str:
        """Fallback PDF processing with PyPDF2."""
        text_content = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text_content += page.extract_text() + "\n"
        return text_content
    
    def _process_pdf_with_ocr(self, file_path: Path) -> str:
        """Process PDF using OCR for scanned documents."""
        try:
            text_content = ""
            pdf_document = fitz.open(file_path)
            
            for page_num in range(pdf_document.page_count):
                page = pdf_document[page_num]
                
                # Convert page to image
                pix = page.get_pixmap()
                img_data = pix.tobytes("ppm")
                img = Image.open(io.BytesIO(img_data))
                
                # OCR the image
                page_text = pytesseract.image_to_string(img)
                text_content += page_text + "\n"
            
            pdf_document.close()
            return text_content
            
        except Exception as e:
            logger.error(f"OCR processing failed: {e}")
            return ""
    
    def _process_docx(self, file_path: Path) -> str:
        """Extract text from DOCX files."""
        try:
            doc = Document(file_path)
            text_content = ""
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_content += cell.text + " "
                    text_content += "\n"
            
            return text_content
            
        except Exception as e:
            logger.error(f"Error processing DOCX {file_path}: {e}")
            return ""
    
    def _process_txt(self, file_path: Path) -> str:
        """Process plain text files."""
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        return file.read()
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, read as binary and decode with errors='ignore'
            with open(file_path, 'rb') as file:
                return file.read().decode('utf-8', errors='ignore')
                
        except Exception as e:
            logger.error(f"Error processing TXT {file_path}: {e}")
            return ""
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text."""
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters but keep punctuation
        text = re.sub(r'[^\w\s\.\,\!\?\;\:\-\(\)]', ' ', text)
        
        # Remove multiple spaces
        text = re.sub(r' +', ' ', text)
        
        # Strip leading/trailing whitespace
        text = text.strip()
        
        return text
    
    def _get_text_statistics(self, text: str) -> Dict:
        """Generate basic statistics about the extracted text."""
        if not text:
            return {}
        
        # Basic counts
        char_count = len(text)
        word_count = len(text.split())
        sentence_count = len([s for s in text.split('.') if s.strip()])
        paragraph_count = len([p for p in text.split('\n') if p.strip()])
        
        # Language detection
        try:
            blob = TextBlob(text[:1000])  # Use first 1000 chars for language detection
            language = blob.detect_language()
        except:
            language = 'unknown'
        
        # Reading time estimation (average 200 words per minute)
        reading_time_minutes = max(1, round(word_count / 200))
        
        return {
            'character_count': char_count,
            'word_count': word_count,
            'sentence_count': sentence_count,
            'paragraph_count': paragraph_count,
            'detected_language': language,
            'estimated_reading_time_minutes': reading_time_minutes
        }
    
    def batch_process(self, directory_path: str) -> List[Dict]:
        """
        Process multiple documents in a directory.
        
        Args:
            directory_path: Path to directory containing documents
            
        Returns:
            List of processing results for each document
        """
        if not os.path.exists(directory_path):
            raise FileNotFoundError(f"Directory not found: {directory_path}")
        
        directory = Path(directory_path)
        results = []
        
        # Find all supported files
        for file_path in directory.iterdir():
            if file_path.is_file() and file_path.suffix.lower() in self.supported_formats:
                logger.info(f"Processing: {file_path.name}")
                result = self.process_document(str(file_path))
                results.append(result)
        
        return results


# Example usage and testing
if __name__ == "__main__":
    # Initialize processor
    processor = DocumentProcessor()
    
    # Test with a sample file (you'll need to provide actual files)
    # result = processor.process_document("sample_document.pdf")
    # print(f"Processing status: {result['processing_status']}")
    # print(f"Word count: {result['text_statistics'].get('word_count', 0)}")
    # print(f"First 200 characters: {result['processed_text'][:200]}...")
    
    print("Document processor initialized successfully!")
    print(f"Supported formats: {processor.supported_formats}")