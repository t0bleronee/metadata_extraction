
"""
Streamlit Web Application for Automated Metadata Generation
Main application file that integrates all modules into a user-friendly interface
"""

import streamlit as st
import pandas as pd
import json
import io
import time
import tempfile
import shutil
from pathlib import Path
import plotly.express as px
import plotly.graph_objects as go
from wordcloud import WordCloud
import matplotlib.pyplot as plt

# Import your actual processing modules from the src folder
import sys
from pathlib import Path

# Add src directory to Python path
src_path = Path(__file__).parent
sys.path.append(str(src_path))


import os
import io
import logging
from typing import Dict, List, Optional, Tuple
from pathlib import Path

# Document processing libraries
import PyPDF2
import pdfplumber
from docx import Document
import pytesseract
from PIL import Image
import fitz  # PyMuPDF for better PDF handling

# Text processing
import re
from textblob import TextBlob

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DocumentProcessor:
    """
    A comprehensive document processor that extracts text from various formats
    and prepares it for metadata generation.
    """
    
    def __init__(self, tesseract_path: Optional[str] = None):
        """
        Initialize the document processor.
        
        Args:
            tesseract_path: Path to tesseract executable (if not in PATH)
        """
        self.supported_formats = ['.pdf', '.docx', '.txt', '.doc']
        
        # Set tesseract path if provided
        if tesseract_path:
            pytesseract.pytesseract.tesseract_cmd = tesseract_path
        
        # Test OCR availability
        self._test_ocr()
    
    def _test_ocr(self) -> bool:
        """Test if OCR is working properly."""
        try:
            pytesseract.get_tesseract_version()
            logger.info("OCR (Tesseract) is available")
            return True
        except Exception as e:
            logger.warning(f"OCR not available: {e}")
            return False
    
    def process_document(self, file_path: str) -> Dict:
        """
        Process a document and extract text content with metadata.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Dictionary containing extracted text and basic metadata
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_path = Path(file_path)
        file_extension = file_path.suffix.lower()
        
        if file_extension not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_extension}")
        
        # Basic file metadata
        file_stats = file_path.stat()
        basic_metadata = {
            'filename': file_path.name,
            'file_size': file_stats.st_size,
            'file_extension': file_extension,
            'file_path': str(file_path)
        }
        
        # Extract text based on file type
        try:
            if file_extension == '.pdf':
                text_content = self._process_pdf(file_path)
            elif file_extension in ['.docx', '.doc']:
                text_content = self._process_docx(file_path)
            elif file_extension == '.txt':
                text_content = self._process_txt(file_path)
            else:
                raise ValueError(f"Handler not implemented for {file_extension}")
            
            # Process and clean text
            processed_text = self._clean_text(text_content)
            
            # Basic text statistics
            text_stats = self._get_text_statistics(processed_text)
            
            return {
                'raw_text': text_content,
                'processed_text': processed_text,
                'file_metadata': basic_metadata,
                'text_statistics': text_stats,
                'processing_status': 'success'
            }
            
        except Exception as e:
            logger.error(f"Error processing {file_path}: {e}")
            return {
                'raw_text': '',
                'processed_text': '',
                'file_metadata': basic_metadata,
                'text_statistics': {},
                'processing_status': 'failed',
                'error': str(e)
            }
    
    def _process_pdf(self, file_path: Path) -> str:
        """Extract text from PDF files with OCR fallback."""
        text_content = ""
        
        try:
            # Try pdfplumber first (better for structured PDFs)
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_content += page_text + "\n"
            
            # If we got very little text, try OCR
            if len(text_content.strip()) < 100:
                logger.info(f"Low text extraction from {file_path.name}, trying OCR...")
                text_content = self._process_pdf_with_ocr(file_path)
                
        except Exception as e:
            logger.warning(f"pdfplumber failed for {file_path.name}: {e}")
            # Fallback to PyPDF2
            try:
                text_content = self._process_pdf_pypdf2(file_path)
            except Exception as e2:
                logger.warning(f"PyPDF2 also failed: {e2}")
                # Last resort: OCR
                text_content = self._process_pdf_with_ocr(file_path)
        
        return text_content
    
    def _process_pdf_pypdf2(self, file_path: Path) -> str:
        """Fallback PDF processing with PyPDF2."""
        text_content = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text_content += page.extract_text() + "\n"
        return text_content
    
    def _process_pdf_with_ocr(self, file_path: Path) -> str:
        """Process PDF using OCR for scanned documents."""
        try:
            text_content = ""
            pdf_document = fitz.open(file_path)
            
            for page_num in range(pdf_document.page_count):
                page = pdf_document[page_num]
                
                # Convert page to image
                pix = page.get_pixmap()
                img_data = pix.tobytes("ppm")
                img = Image.open(io.BytesIO(img_data))
                
                # OCR the image
                page_text = pytesseract.image_to_string(img)
                text_content += page_text + "\n"
            
            pdf_document.close()
            return text_content
            
        except Exception as e:
            logger.error(f"OCR processing failed: {e}")
            return ""
    
    def _process_docx(self, file_path: Path) -> str:
        """Extract text from DOCX files."""
        try:
            doc = Document(file_path)
            text_content = ""
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_content += cell.text + " "
                    text_content += "\n"
            
            return text_content
            
        except Exception as e:
            logger.error(f"Error processing DOCX {file_path}: {e}")
            return ""
    
    def _process_txt(self, file_path: Path) -> str:
        """Process plain text files."""
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        return file.read()
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, read as binary and decode with errors='ignore'
            with open(file_path, 'rb') as file:
                return file.read().decode('utf-8', errors='ignore')
                
        except Exception as e:
            logger.error(f"Error processing TXT {file_path}: {e}")
            return ""
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text."""
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters but keep punctuation
        text = re.sub(r'[^\w\s\.\,\!\?\;\:\-\(\)]', ' ', text)
        
        # Remove multiple spaces
        text = re.sub(r' +', ' ', text)
        
        # Strip leading/trailing whitespace
        text = text.strip()
        
        return text
    
    def _get_text_statistics(self, text: str) -> Dict:
        """Generate basic statistics about the extracted text."""
        if not text:
            return {}
        
        # Basic counts
        char_count = len(text)
        word_count = len(text.split())
        sentence_count = len([s for s in text.split('.') if s.strip()])
        paragraph_count = len([p for p in text.split('\n') if p.strip()])
        
        # Language detection
        try:
            blob = TextBlob(text[:1000])  # Use first 1000 chars for language detection
            language = blob.detect_language()
        except:
            language = 'unknown'
        
        # Reading time estimation (average 200 words per minute)
        reading_time_minutes = max(1, round(word_count / 200))
        
        return {
            'character_count': char_count,
            'word_count': word_count,
            'sentence_count': sentence_count,
            'paragraph_count': paragraph_count,
            'detected_language': language,
            'estimated_reading_time_minutes': reading_time_minutes
        }
    
    def batch_process(self, directory_path: str) -> List[Dict]:
        """
        Process multiple documents in a directory.
        
        Args:
            directory_path: Path to directory containing documents
            
        Returns:
            List of processing results for each document
        """
        if not os.path.exists(directory_path):
            raise FileNotFoundError(f"Directory not found: {directory_path}")
        
        directory = Path(directory_path)
        results = []
        
        # Find all supported files
        for file_path in directory.iterdir():
            if file_path.is_file() and file_path.suffix.lower() in self.supported_formats:
                logger.info(f"Processing: {file_path.name}")
                result = self.process_document(str(file_path))
                results.append(result)
        
        return results


# Example usage and testing
#if __name__ == "__main__":
    # Initialize processor
   # processor = DocumentProcessor()
    
    # Test with a sample file (you'll need to provide actual files)
    # result = processor.process_document("sample_document.pdf")
    # print(f"Processing status: {result['processing_status']}")
    # print(f"Word count: {result['text_statistics'].get('word_count', 0)}")
    # print(f"First 200 characters: {result['processed_text'][:200]}...")
    
   # print("Document processor initialized successfully!")
   # print(f"Supported formats: {processor.supported_formats}")

import re
import logging
from collections import Counter, defaultdict
from typing import Dict, List, Tuple, Set, Optional
from dataclasses import dataclass
import unicodedata

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

@dataclass
class EntityInfo:
    """Information about extracted entities"""
    text: str
    entity_type: str
    confidence: float = 1.0
    context: str = ""

@dataclass
class TopicInfo:
    """Information about extracted topics"""
    topic: str
    keywords: List[str]
    confidence: float
    sentences: List[str]

class NLPAnalyzer:
    """
    Advanced NLP analysis for document metadata generation
    Uses rule-based and statistical approaches for robust analysis
    """
    
    def __init__(self):
        """Initialize the NLP analyzer with patterns and vocabularies"""
        logger.info("Initializing NLP Analyzer...")
        
        # Entity recognition patterns
        self.entity_patterns = {
            'email': re.compile(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'),
            'phone': re.compile(r'(\+\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'),
            'url': re.compile(r'https?://(?:[-\w.])+(?:\:[0-9]+)?(?:/(?:[\w/_.])*(?:\?(?:[\w&=%.])*)?(?:\#(?:[\w.])*)?)?'),
            'date': re.compile(r'\b(?:\d{1,2}/\d{1,2}/\d{2,4}|\d{4}-\d{2}-\d{2}|(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]* \d{1,2},? \d{4})\b', re.IGNORECASE),
            'currency': re.compile(r'\$\d{1,3}(?:,\d{3})*(?:\.\d{2})?|\b\d{1,3}(?:,\d{3})*(?:\.\d{2})? (?:USD|EUR|GBP|CAD)\b'),
            'percentage': re.compile(r'\b\d+(?:\.\d+)?%\b'),
            'organization': re.compile(r'\b(?:Inc|LLC|Corp|Ltd|Company|Organization|University|Institute|Department|Agency|Foundation|Association)\b', re.IGNORECASE),
        }
        
        # Common stop words for topic extraction
        self.stop_words = {
            'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with', 'by',
            'is', 'are', 'was', 'were', 'be', 'been', 'being', 'have', 'has', 'had', 'do', 'does', 'did',
            'will', 'would', 'could', 'should', 'may', 'might', 'can', 'this', 'that', 'these', 'those',
            'i', 'you', 'he', 'she', 'it', 'we', 'they', 'me', 'him', 'her', 'us', 'them', 'my', 'your',
            'his', 'her', 'its', 'our', 'their', 'myself', 'yourself', 'himself', 'herself', 'itself',
            'ourselves', 'yourselves', 'themselves', 'what', 'which', 'who', 'whom', 'whose', 'where',
            'when', 'why', 'how', 'all', 'any', 'both', 'each', 'few', 'more', 'most', 'other', 'some',
            'such', 'no', 'nor', 'not', 'only', 'own', 'same', 'so', 'than', 'too', 'very', 'can', 'just'
        }
        
        # Technical/domain keywords that indicate topics
        self.domain_keywords = {
            'technology': ['software', 'hardware', 'computer', 'system', 'application', 'program', 'code', 'algorithm', 'data', 'database', 'network', 'server', 'cloud', 'api', 'framework', 'platform'],
            'business': ['revenue', 'profit', 'sales', 'marketing', 'customer', 'client', 'market', 'strategy', 'business', 'company', 'organization', 'management', 'finance', 'budget', 'cost'],
            'research': ['study', 'research', 'analysis', 'experiment', 'hypothesis', 'methodology', 'results', 'conclusion', 'findings', 'data', 'sample', 'statistical', 'significant'],
            'medical': ['patient', 'treatment', 'diagnosis', 'medical', 'health', 'clinical', 'therapy', 'disease', 'symptom', 'medicine', 'doctor', 'hospital', 'healthcare'],
            'legal': ['contract', 'agreement', 'legal', 'law', 'court', 'case', 'plaintiff', 'defendant', 'attorney', 'lawyer', 'litigation', 'settlement', 'regulation', 'compliance'],
            'education': ['student', 'teacher', 'education', 'learning', 'course', 'curriculum', 'school', 'university', 'academic', 'degree', 'diploma', 'graduate', 'undergraduate'],
            'finance': ['investment', 'portfolio', 'stock', 'bond', 'asset', 'liability', 'equity', 'dividend', 'interest', 'loan', 'credit', 'debt', 'financial', 'banking']
        }
        
        logger.info("✅ NLP Analyzer initialized!")
    
    def analyze_document(self, text: str, filename: str = "") -> Dict:
        """
        Perform comprehensive NLP analysis on document text
        
        Args:
            text: Document text to analyze
            filename: Optional filename for context
            
        Returns:
            Dict containing all analysis results
        """
        logger.info(f"Starting NLP analysis for: {filename or 'text input'}")
        
        # Clean and preprocess text
        clean_text = self._preprocess_text(text)
        sentences = self._split_sentences(clean_text)
        words = self._extract_words(clean_text)
        
        # Perform various analyses
        analysis_results = {
            'entities': self._extract_entities(text),
            'topics': self._extract_topics(clean_text, sentences, words),
            'summary': self._generate_summary(sentences),
            'keywords': self._extract_keywords(words),
            'document_structure': self._analyze_structure(text),
            'sentiment': self._analyze_sentiment(clean_text),
            'readability': self._analyze_readability(sentences, words),
            'language_features': self._analyze_language_features(text, words)
        }
        
        logger.info("✅ NLP analysis completed")
        return analysis_results
    
    def _preprocess_text(self, text: str) -> str:
        """Clean and normalize text for analysis"""
        # Normalize unicode characters
        text = unicodedata.normalize('NFKD', text)
        
        # Remove excessive whitespace but preserve paragraph structure
        text = re.sub(r'[ \t]+', ' ', text)
        text = re.sub(r'\n\s*\n', '\n\n', text)
        
        return text.strip()
    
    def _split_sentences(self, text: str) -> List[str]:
        """Split text into sentences using rule-based approach"""
        # Simple sentence splitting - handles most cases
        sentences = re.split(r'(?<=[.!?])\s+(?=[A-Z])', text)
        return [s.strip() for s in sentences if s.strip()]
    
    def _extract_words(self, text: str) -> List[str]:
        """Extract and clean words from text"""
        words = re.findall(r'\b[a-zA-Z]+\b', text.lower())
        return [word for word in words if len(word) > 2]
    
    def _extract_entities(self, text: str) -> Dict[str, List[EntityInfo]]:
        """Extract named entities using pattern matching"""
        entities = defaultdict(list)
        
        for entity_type, pattern in self.entity_patterns.items():
            matches = pattern.finditer(text)
            for match in matches:
                entity_text = match.group().strip()
                if entity_text:  # Avoid empty matches
                    # Get context (20 characters before and after)
                    start = max(0, match.start() - 20)
                    end = min(len(text), match.end() + 20)
                    context = text[start:end].replace('\n', ' ')
                    
                    entities[entity_type].append(EntityInfo(
                        text=entity_text,
                        entity_type=entity_type,
                        confidence=0.9,  # High confidence for pattern matches
                        context=context
                    ))
        
        # Remove duplicates
        for entity_type in entities:
            seen = set()
            unique_entities = []
            for entity in entities[entity_type]:
                if entity.text.lower() not in seen:
                    seen.add(entity.text.lower())
                    unique_entities.append(entity)
            entities[entity_type] = unique_entities
        
        return dict(entities)
    
    def _extract_topics(self, text: str, sentences: List[str], words: List[str]) -> List[TopicInfo]:
        """Extract main topics using keyword analysis and domain detection"""
        topics = []
        
        # Count word frequencies (excluding stop words)
        word_freq = Counter([word for word in words if word not in self.stop_words])
        
        # Identify domain-specific topics
        domain_scores = defaultdict(int)
        domain_keywords_found = defaultdict(list)
        
        for domain, keywords in self.domain_keywords.items():
            for keyword in keywords:
                if keyword in word_freq:
                    domain_scores[domain] += word_freq[keyword]
                    domain_keywords_found[domain].append(keyword)
        
        # Create topic info for top domains
        for domain, score in sorted(domain_scores.items(), key=lambda x: x[1], reverse=True)[:3]:
            if score > 2:  # Minimum threshold
                relevant_sentences = []
                for sentence in sentences[:10]:  # Check first 10 sentences
                    if any(keyword in sentence.lower() for keyword in domain_keywords_found[domain]):
                        relevant_sentences.append(sentence)
                
                topics.append(TopicInfo(
                    topic=domain.title(),
                    keywords=domain_keywords_found[domain][:5],
                    confidence=min(score / 10, 1.0),
                    sentences=relevant_sentences[:3]
                ))
        
        # Add general high-frequency topics
        top_words = [word for word, count in word_freq.most_common(10) if count > 2]
        if top_words:
            topics.append(TopicInfo(
                topic="General",
                keywords=top_words[:5],
                confidence=0.5,
                sentences=sentences[:2]
            ))
        
        return topics
    
    def _extract_keywords(self, words: List[str]) -> List[Tuple[str, int]]:
        """Extract important keywords with frequency"""
        word_freq = Counter([word for word in words if word not in self.stop_words])
        
        # Filter out very common and very rare words
        total_words = len(words)
        keywords = []
        
        for word, freq in word_freq.most_common(20):
            # Skip if too rare (< 0.1%) or too common (> 5%)
            percentage = freq / total_words
            if 0.001 <= percentage <= 0.05 and len(word) > 3:
                keywords.append((word, freq))
        
        return keywords
    
    def _generate_summary(self, sentences: List[str], max_sentences: int = 3) -> str:
        """Generate extractive summary using sentence scoring"""
        if not sentences:
            return ""
        
        if len(sentences) <= max_sentences:
            return " ".join(sentences)
        
        # Score sentences based on length and position
        sentence_scores = []
        for i, sentence in enumerate(sentences):
            score = 0
            
            # Position score (earlier sentences are more important)
            position_score = 1.0 - (i / len(sentences))
            score += position_score * 0.3
            
            # Length score (medium length sentences preferred)
            word_count = len(sentence.split())
            if 10 <= word_count <= 30:
                score += 0.4
            elif 5 <= word_count < 10 or 30 < word_count <= 50:
                score += 0.2
            
            # Keyword score - sentences with important words
            important_words = ['important', 'significant', 'key', 'main', 'primary', 'conclusion', 'result']
            if any(word in sentence.lower() for word in important_words):
                score += 0.3
            
            sentence_scores.append((sentence, score))
        
        # Select top sentences
        top_sentences = sorted(sentence_scores, key=lambda x: x[1], reverse=True)[:max_sentences]
        
        # Return in original order
        selected_indices = []
        for sentence, _ in top_sentences:
            try:
                idx = sentences.index(sentence)
                selected_indices.append(idx)
            except ValueError:
                continue
        
        selected_indices.sort()
        summary_sentences = [sentences[i] for i in selected_indices]
        
        return " ".join(summary_sentences)
    
    def _analyze_structure(self, text: str) -> Dict:
        """Analyze document structure"""
        structure = {
            'has_title': False,
            'has_headings': False,
            'has_lists': False,
            'has_tables': False,
            'sections': 0,
            'paragraphs': len(text.split('\n\n'))
        }
        
        # Check for title (first line in caps or with specific patterns)
        lines = text.split('\n')
        if lines:
            first_line = lines[0].strip()
            if (first_line.isupper() or 
                any(char in first_line for char in [':', '=', '-']) and len(first_line) < 100):
                structure['has_title'] = True
        
        # Check for headings (lines with specific patterns)
        heading_patterns = [
            r'^[A-Z][A-Za-z\s]+:$',  # "Section Name:"
            r'^#+\s',                 # Markdown headers
            r'^[IVX]+\.',            # Roman numerals
            r'^\d+\.',               # Numbered sections
        ]
        
        for line in lines:
            line = line.strip()
            if any(re.match(pattern, line) for pattern in heading_patterns):
                structure['has_headings'] = True
                structure['sections'] += 1
        
        # Check for lists
        if re.search(r'^\s*[-*•]\s', text, re.MULTILINE):
            structure['has_lists'] = True
        
        # Check for tables (simple heuristic)
        if '|' in text or '\t' in text:
            structure['has_tables'] = True
        
        return structure
    
    def _analyze_sentiment(self, text: str) -> Dict:
        """Basic sentiment analysis using word lists"""
        positive_words = {
            'good', 'great', 'excellent', 'amazing', 'wonderful', 'fantastic', 'outstanding',
            'positive', 'success', 'successful', 'achieve', 'achievement', 'improve', 'improvement',
            'benefit', 'advantage', 'effective', 'efficient', 'valuable', 'important', 'significant'
        }
        
        negative_words = {
            'bad', 'terrible', 'awful', 'horrible', 'poor', 'negative', 'fail', 'failure',
            'problem', 'issue', 'challenge', 'difficult', 'hard', 'impossible', 'wrong',
            'error', 'mistake', 'concern', 'risk', 'threat', 'danger', 'loss', 'decline'
        }
        
        words = re.findall(r'\b[a-zA-Z]+\b', text.lower())
        
        positive_count = sum(1 for word in words if word in positive_words)
        negative_count = sum(1 for word in words if word in negative_words)
        total_sentiment_words = positive_count + negative_count
        
        if total_sentiment_words == 0:
            sentiment = "neutral"
            score = 0.0
        else:
            score = (positive_count - negative_count) / total_sentiment_words
            if score > 0.1:
                sentiment = "positive"
            elif score < -0.1:
                sentiment = "negative"
            else:
                sentiment = "neutral"
        
        return {
            'sentiment': sentiment,
            'score': score,
            'positive_words': positive_count,
            'negative_words': negative_count
        }
    
    def _analyze_readability(self, sentences: List[str], words: List[str]) -> Dict:
        """Calculate readability metrics"""
        if not sentences or not words:
            return {'flesch_score': 0, 'reading_level': 'unknown'}
        
        avg_sentence_length = len(words) / len(sentences)
        
        # Count syllables (rough approximation)
        def count_syllables(word):
            vowels = 'aeiouy'
            syllables = 0
            prev_char_vowel = False
            
            for char in word.lower():
                if char in vowels:
                    if not prev_char_vowel:
                        syllables += 1
                    prev_char_vowel = True
                else:
                    prev_char_vowel = False
            
            return max(1, syllables)  # Every word has at least 1 syllable
        
        total_syllables = sum(count_syllables(word) for word in words)
        avg_syllables_per_word = total_syllables / len(words)
        
        # Flesch Reading Ease Score
        flesch_score = 206.835 - (1.015 * avg_sentence_length) - (84.6 * avg_syllables_per_word)
        
        # Determine reading level
        if flesch_score >= 90:
            reading_level = "very_easy"
        elif flesch_score >= 80:
            reading_level = "easy"
        elif flesch_score >= 70:
            reading_level = "fairly_easy"
        elif flesch_score >= 60:
            reading_level = "standard"
        elif flesch_score >= 50:
            reading_level = "fairly_difficult"
        elif flesch_score >= 30:
            reading_level = "difficult"
        else:
            reading_level = "very_difficult"
        
        return {
            'flesch_score': round(flesch_score, 1),
            'reading_level': reading_level,
            'avg_sentence_length': round(avg_sentence_length, 1),
            'avg_syllables_per_word': round(avg_syllables_per_word, 2)
        }
    
    def _analyze_language_features(self, text: str, words: List[str]) -> Dict:
        """Analyze various language features"""
        features = {
            'lexical_diversity': 0,
            'formality_score': 0,
            'complexity_indicators': []
        }
        
        # Lexical diversity (unique words / total words)
        if words:
            unique_words = set(words)
            features['lexical_diversity'] = round(len(unique_words) / len(words), 3)
        
        # Formality indicators
        formal_indicators = ['therefore', 'however', 'furthermore', 'moreover', 'consequently', 
                           'nevertheless', 'accordingly', 'thus', 'hence', 'whereas']
        informal_indicators = ['really', 'pretty', 'quite', 'very', 'totally', 'basically', 
                             'actually', 'literally', 'obviously', 'definitely']
        
        formal_count = sum(1 for word in words if word in formal_indicators)
        informal_count = sum(1 for word in words if word in informal_indicators)
        
        if formal_count + informal_count > 0:
            features['formality_score'] = formal_count / (formal_count + informal_count)
        
        # Complexity indicators
        if re.search(r'\b(?:complex|complicated|sophisticated|intricate)\b', text, re.IGNORECASE):
            features['complexity_indicators'].append('complex_vocabulary')
        
        if re.search(r'[;:]', text):
            features['complexity_indicators'].append('complex_punctuation')
        
        # Average word length
        if words:
            avg_word_length = sum(len(word) for word in words) / len(words)
            if avg_word_length > 5:
                features['complexity_indicators'].append('long_words')
        
        return features

# Example usage and testing
#if __name__ == "__main__":
    # Test the analyzer
#    analyzer = NLPAnalyzer()
# sample_text = """
#     Artificial Intelligence and Machine Learning Report
#     
#     Introduction
#     This report examines the current state of artificial intelligence and machine learning 
#     technologies in business applications. Our research shows significant improvements in 
#     automated decision-making systems.
#     
#     Key Findings
#     - Machine learning algorithms have improved accuracy by 25%
#     - Implementation costs have decreased by $50,000 annually
#     - Customer satisfaction increased to 92%
#     
#     Methodology
#     We analyzed data from 150 companies over 12 months. The study included both quantitative 
#     metrics and qualitative feedback from users.
#     
#     Conclusion
#     The results demonstrate that AI technologies provide substantial benefits for modern 
#     businesses. We recommend continued investment in these systems.
#     
#     Contact: research@company.com
#     Phone: (555) 123-4567
#     """
# 
# results = analyzer.analyze_document(sample_text, "ai_report.txt")
# 
# print("🔍 NLP Analysis Results:")
# print("=" * 50)
# 
# print("\n📊 Entities Found:")
# for entity_type, entities in results['entities'].items():
#     if entities:
#         print(f"  {entity_type.title()}: {[e.text for e in entities]}")
# 
# print(f"\n🏷️ Topics ({len(results['topics'])}):")
# for topic in results['topics']:
#     print(f"  - {topic.topic} (confidence: {topic.confidence:.2f})")
#     print(f"    Keywords: {', '.join(topic.keywords[:3])}")
# 
# print(f"\n📝 Summary:")
# print(f"  {results['summary']}")
# 
# print(f"\n🔤 Top Keywords:")
# for keyword, freq in results['keywords'][:5]:
#     print(f"  - {keyword}: {freq}")
# 
# print(f"\n📖 Readability:")
# print(f"  - Reading level: {results['readability']['reading_level']}")
# print(f"  - Flesch score: {results['readability']['flesch_score']}")
# 
# print(f"\n💭 Sentiment: {results['sentiment']['sentiment']} ({results['sentiment']['score']:.2f})")



import json
import uuid
from datetime import datetime
from typing import Dict, List, Any, Optional
from dataclasses import dataclass, asdict
from pathlib import Path
import logging

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

@dataclass
class DocumentMetadata:
    """Structured metadata for a document"""
    # Basic document information
    document_id: str
    filename: str
    file_size: int
    file_type: str
    processing_date: str
    
    # Content analysis
    title: Optional[str]
    summary: str
    language: str
    
    # Text statistics
    word_count: int
    sentence_count: int
    paragraph_count: int
    reading_time_minutes: int
    
    # Topics and keywords
    primary_topics: List[str]
    keywords: List[str]
    entities: Dict[str, List[str]]
    
    # Document characteristics
    document_type: str
    formality_level: str
    complexity_level: str
    sentiment: str
    
    # Technical metadata
    readability_score: float
    readability_level: str
    structure_score: float
    
    # Additional fields
    tags: List[str]
    categories: List[str]
    confidence_score: float

class MetadataGenerator:
    """
    Generates structured metadata from document processing and NLP analysis results
    """
    
    def __init__(self):
        """Initialize the metadata generator"""
        logger.info("Initializing Metadata Generator...")
        
        # Document type classification patterns
        self.document_type_patterns = {
            'research_paper': ['abstract', 'methodology', 'conclusion', 'references', 'hypothesis'],
            'business_report': ['executive summary', 'revenue', 'quarterly', 'analysis', 'recommendations'],
            'technical_document': ['system', 'architecture', 'implementation', 'configuration', 'api'],
            'legal_document': ['agreement', 'contract', 'terms', 'conditions', 'liability'],
            'manual': ['instructions', 'steps', 'procedure', 'guide', 'how to'],
            'presentation': ['slide', 'agenda', 'overview', 'introduction', 'thank you'],
            'article': ['article', 'author', 'published', 'journal', 'volume'],
            'proposal': ['proposal', 'budget', 'timeline', 'objectives', 'deliverables']
        }
        
        # Category mapping based on topics
        self.category_mapping = {
            'technology': ['Information Technology', 'Software Development', 'Digital Innovation'],
            'business': ['Business Strategy', 'Finance', 'Marketing', 'Management'],
            'research': ['Academic Research', 'Scientific Study', 'Data Analysis'],
            'medical': ['Healthcare', 'Medical Research', 'Clinical Studies'],
            'legal': ['Legal Documents', 'Contracts', 'Compliance'],
            'education': ['Educational Content', 'Training Materials', 'Academic'],
            'finance': ['Financial Analysis', 'Investment', 'Banking']
        }
        
        logger.info("✅ Metadata Generator initialized!")
    
    def generate_metadata(self, 
                         document_result: Dict, 
                         nlp_result: Dict,
                         custom_tags: List[str] = None) -> DocumentMetadata:
        """
        Generate comprehensive metadata from processing results
        
        Args:
            document_result: Output from DocumentProcessor
            nlp_result: Output from NLPAnalyzer
            custom_tags: Optional custom tags to add
            
        Returns:
            DocumentMetadata object with all extracted information
        """
        logger.info(f"Generating metadata for: {document_result.get('file_metadata', {}).get('filename', 'unknown')}")
        
        # Extract basic information
        file_meta = document_result.get('file_metadata', {})
        text_stats = document_result.get('text_statistics', {})
        
        # Generate unique document ID
        doc_id = str(uuid.uuid4())[:8]
        
        # Extract title (heuristic approach)
        title = self._extract_title(document_result.get('processed_text', ''))
        
        # Determine document type
        doc_type = self._classify_document_type(
            document_result.get('processed_text', ''), 
            nlp_result
        )
        
        # Extract and process topics
        topics = self._process_topics(nlp_result.get('topics', []))
        
        # Extract keywords
        keywords = self._process_keywords(nlp_result.get('keywords', []))
        
        # Process entities
        entities = self._process_entities(nlp_result.get('entities', {}))
        
        # Determine categories
        categories = self._determine_categories(topics, entities, keywords)
        
        # Generate tags
        tags = self._generate_tags(topics, keywords, doc_type, custom_tags or [])
        
        # Calculate scores
        confidence_score = self._calculate_confidence_score(nlp_result)
        structure_score = self._calculate_structure_score(nlp_result.get('document_structure', {}))
        
        # Determine complexity and formality
        complexity_level = self._determine_complexity_level(nlp_result)
        formality_level = self._determine_formality_level(nlp_result)
        
        # Create metadata object
        metadata = DocumentMetadata(
            # Basic information
            document_id=doc_id,
            filename=file_meta.get('filename', 'unknown'),
            file_size=file_meta.get('file_size', 0),
            file_type=file_meta.get('file_extension', '').upper().replace('.', ''),
            processing_date=datetime.now().isoformat(),
            
            # Content
            title=title,
            summary=nlp_result.get('summary', ''),
            language=text_stats.get('detected_language', 'unknown'),
            
            # Statistics
            word_count=text_stats.get('word_count', 0),
            sentence_count=text_stats.get('sentence_count', 0),
            paragraph_count=text_stats.get('paragraph_count', 0),
            reading_time_minutes=text_stats.get('estimated_reading_time_minutes', 0),
            
            # Analysis results
            primary_topics=topics[:5],  # Top 5 topics
            keywords=keywords[:10],     # Top 10 keywords
            entities=entities,
            
            # Classification
            document_type=doc_type,
            formality_level=formality_level,
            complexity_level=complexity_level,
            sentiment=nlp_result.get('sentiment', {}).get('sentiment', 'neutral'),
            
            # Technical scores
            readability_score=nlp_result.get('readability', {}).get('flesch_score', 0),
            readability_level=nlp_result.get('readability', {}).get('reading_level', 'unknown'),
            structure_score=structure_score,
            
            # Organization
            tags=tags,
            categories=categories,
            confidence_score=confidence_score
        )
        
        logger.info("✅ Metadata generation completed")
        return metadata
    
    def _extract_title(self, text: str) -> Optional[str]:
        """Extract document title using heuristics"""
        if not text:
            return None
        
        lines = text.split('\n')
        for line in lines[:5]:  # Check first 5 lines
            line = line.strip()
            if line and len(line) > 5 and len(line) < 100:
                # Check if it looks like a title
                if (line.isupper() or 
                    line.count(' ') < 10 or 
                    any(char in line for char in [':', '=', '-']) or
                    not line.endswith('.')):
                    return line
        
        # Fallback: use first sentence if it's short enough
        sentences = text.split('.')
        if sentences and len(sentences[0]) < 100:
            return sentences[0].strip()
        
        return None
    
    def _classify_document_type(self, text: str, nlp_result: Dict) -> str:
        """Classify document type based on content patterns"""
        text_lower = text.lower()
        
        # Score each document type
        type_scores = {}
        for doc_type, patterns in self.document_type_patterns.items():
            score = sum(1 for pattern in patterns if pattern in text_lower)
            if score > 0:
                type_scores[doc_type] = score
        
        # Also consider structure
        structure = nlp_result.get('document_structure', {})
        if structure.get('has_title') and structure.get('has_headings'):
            type_scores['report'] = type_scores.get('report', 0) + 2
        
        # Return highest scoring type or 'document' as default
        if type_scores:
            return max(type_scores.items(), key=lambda x: x[1])[0].replace('_', ' ').title()
        
        return 'Document'
    
    def _process_topics(self, topics: List) -> List[str]:
        """Process and clean topic information"""
        processed_topics = []
        for topic in topics:
            if hasattr(topic, 'topic'):
                topic_name = topic.topic
            elif isinstance(topic, dict):
                topic_name = topic.get('topic', '')
            else:
                topic_name = str(topic)
            
            if topic_name and topic_name.lower() != 'general':
                processed_topics.append(topic_name.title())
        
        return processed_topics
    
    def _process_keywords(self, keywords: List) -> List[str]:
        """Process and clean keyword information"""
        processed_keywords = []
        for keyword in keywords:
            if isinstance(keyword, tuple):
                word = keyword[0]
            elif isinstance(keyword, dict):
                word = keyword.get('word', '')
            else:
                word = str(keyword)
            
            if word and len(word) > 3:
                processed_keywords.append(word.lower())
        
        return processed_keywords
    
    def _process_entities(self, entities: Dict) -> Dict[str, List[str]]:
        """Process and clean entity information"""
        processed_entities = {}
        
        for entity_type, entity_list in entities.items():
            clean_entities = []
            for entity in entity_list:
                if hasattr(entity, 'text'):
                    entity_text = entity.text
                elif isinstance(entity, dict):
                    entity_text = entity.get('text', '')
                else:
                    entity_text = str(entity)
                
                if entity_text:
                    clean_entities.append(entity_text)
            
            if clean_entities:
                processed_entities[entity_type.replace('_', ' ').title()] = clean_entities[:5]  # Limit to 5 per type
        
        return processed_entities
    
    def _determine_categories(self, topics: List[str], entities: Dict, keywords: List[str]) -> List[str]:
        """Determine document categories based on content analysis"""
        categories = set()
        
        # Category from topics
        for topic in topics:
            topic_lower = topic.lower()
            for domain, cats in self.category_mapping.items():
                if domain in topic_lower:
                    categories.update(cats[:2])  # Add first 2 categories
        
        # Category from entities
        if 'Email' in entities or 'Phone' in entities:
            categories.add('Contact Information')
        if 'Organization' in entities:
            categories.add('Organizational')
        if 'Date' in entities:
            categories.add('Time-Sensitive')
        
        # Category from keywords
        keyword_str = ' '.join(keywords).lower()
        if any(word in keyword_str for word in ['report', 'analysis', 'study']):
            categories.add('Analytical')
        if any(word in keyword_str for word in ['policy', 'procedure', 'guideline']):
            categories.add('Procedural')
        
        return list(categories)[:5]  # Limit to 5 categories
    
    def _generate_tags(self, topics: List[str], keywords: List[str], doc_type: str, custom_tags: List[str]) -> List[str]:
        """Generate relevant tags for the document"""
        tags = set()
        
        # Add topic-based tags
        tags.update([topic.lower().replace(' ', '_') for topic in topics])
        
        # Add document type tag
        tags.add(doc_type.lower().replace(' ', '_'))
        
        # Add high-frequency keyword tags
        tags.update(keywords[:5])
        
        # Add custom tags
        tags.update([tag.lower().replace(' ', '_') for tag in custom_tags])
        
        # Clean and return
        clean_tags = [tag for tag in tags if len(tag) > 2 and len(tag) < 20]
        return sorted(clean_tags)[:10]  # Limit to 10 tags
    
    def _calculate_confidence_score(self, nlp_result: Dict) -> float:
        """Calculate overall confidence score for the metadata"""
        scores = []
        
        # Topic confidence
        topics = nlp_result.get('topics', [])
        if topics:
            topic_confidences = []
            for topic in topics:
                if hasattr(topic, 'confidence'):
                    topic_confidences.append(topic.confidence)
                elif isinstance(topic, dict):
                    topic_confidences.append(topic.get('confidence', 0.5))
            if topic_confidences:
                scores.append(sum(topic_confidences) / len(topic_confidences))
        
        # Entity confidence (pattern-based entities have high confidence)
        entities = nlp_result.get('entities', {})
        if entities:
            scores.append(0.8)  # High confidence for pattern-based extraction
        
        # Readability confidence
        readability = nlp_result.get('readability', {})
        if readability.get('flesch_score', 0) > 0:
            scores.append(0.7)
        
        # Structure confidence
        structure = nlp_result.get('document_structure', {})
        structure_indicators = sum([
            structure.get('has_title', False),
            structure.get('has_headings', False),
            structure.get('sections', 0) > 0
        ])
        scores.append(structure_indicators / 3)
        
        return round(sum(scores) / len(scores) if scores else 0.5, 2)
    
    def _calculate_structure_score(self, structure: Dict) -> float:
        """Calculate document structure score"""
        score = 0
        
        # Points for structural elements
        if structure.get('has_title'):
            score += 0.2
        if structure.get('has_headings'):
            score += 0.3
        if structure.get('has_lists'):
            score += 0.2
        if structure.get('sections', 0) > 0:
            score += 0.2
        if structure.get('paragraphs', 0) > 1:
            score += 0.1
        
        return round(min(score, 1.0), 2)
    
    def _determine_complexity_level(self, nlp_result: Dict) -> str:
        """Determine document complexity level"""
        readability = nlp_result.get('readability', {})
        language_features = nlp_result.get('language_features', {})
        
        flesch_score = readability.get('flesch_score', 50)
        lexical_diversity = language_features.get('lexical_diversity', 0.5)
        complexity_indicators = len(language_features.get('complexity_indicators', []))
        
        # Calculate complexity score
        complexity_score = 0
        
        if flesch_score < 30:
            complexity_score += 3
        elif flesch_score < 50:
            complexity_score += 2
        elif flesch_score < 70:
            complexity_score += 1
        
        if lexical_diversity > 0.7:
            complexity_score += 2
        elif lexical_diversity > 0.5:
            complexity_score += 1
        
        complexity_score += complexity_indicators
        
        if complexity_score >= 5:
            return 'High'
        elif complexity_score >= 3:
            return 'Medium'
        else:
            return 'Low'
    
    def _determine_formality_level(self, nlp_result: Dict) -> str:
        """Determine document formality level"""
        language_features = nlp_result.get('language_features', {})
        formality_score = language_features.get('formality_score', 0.5)
        
        if formality_score > 0.7:
            return 'Formal'
        elif formality_score > 0.3:
            return 'Semi-formal'
        else:
            return 'Informal'
    
    def export_metadata(self, metadata: DocumentMetadata, format: str = 'json') -> str:
        """
        Export metadata in various formats
        
        Args:
            metadata: DocumentMetadata object
            format: Export format ('json', 'xml', 'yaml')
            
        Returns:
            Formatted metadata string
        """
        metadata_dict = asdict(metadata)
        
        if format.lower() == 'json':
            return json.dumps(metadata_dict, indent=2, ensure_ascii=False)
        
        elif format.lower() == 'xml':
            return self._dict_to_xml(metadata_dict, 'document_metadata')
        
        elif format.lower() == 'yaml':
            try:
                import yaml
                return yaml.dump(metadata_dict, default_flow_style=False)
            except ImportError:
                logger.warning("PyYAML not installed, falling back to JSON")
                return json.dumps(metadata_dict, indent=2)
        
        else:
            return json.dumps(metadata_dict, indent=2)
    
    def _dict_to_xml(self, data: Dict, root_name: str) -> str:
        """Convert dictionary to XML format"""
        def dict_to_xml_recursive(d, root):
            xml_str = f"<{root}>\n"
            for key, value in d.items():
                if isinstance(value, dict):
                    xml_str += f"  {dict_to_xml_recursive(value, key)}\n"
                elif isinstance(value, list):
                    xml_str += f"  <{key}>\n"
                    for item in value:
                        if isinstance(item, str):
                            xml_str += f"    <item>{item}</item>\n"
                        else:
                            xml_str += f"    <item>{str(item)}</item>\n"
                    xml_str += f"  </{key}>\n"
                else:
                    xml_str += f"  <{key}>{str(value)}</{key}>\n"
            xml_str += f"</{root}>"
            return xml_str
        
        return dict_to_xml_recursive(data, root_name)
    
    def save_metadata(self, metadata: DocumentMetadata, output_path: str, format: str = 'json'):
        """Save metadata to file"""
        formatted_data = self.export_metadata(metadata, format)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(formatted_data)
        
        logger.info(f"Metadata saved to: {output_path}")


# Example usage and testing
#if __name__ == "__main__":
    # Test the metadata generator
   # generator = MetadataGenerator()
    
   

# Note: Your modules are located in the src/ folder:
# - src/document_processor.py (contains DocumentProcessor class)
# - src/nlp_analyzer.py (contains NLPAnalyzer class)  
# - src/metadata_generator.py (contains MetadataGenerator class)

# Initialize components
@st.cache_resource
def initialize_components():
    """Initialize all processing components"""
    processor = DocumentProcessor()    # Your real DocumentProcessor class
    analyzer = NLPAnalyzer()          # Your real NLPAnalyzer class  
    generator = MetadataGenerator()    # Your real MetadataGenerator class
    return processor, analyzer, generator

def create_analytics_dashboard(metadata_list):
    """Create analytics dashboard for processed documents"""
    if not metadata_list:
        return
    
    st.subheader("📊 Document Analytics Dashboard")
    
    # Metrics row
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.metric("Total Documents", len(metadata_list))
    
    with col2:
        avg_words = sum(m.word_count for m in metadata_list) / len(metadata_list)
        st.metric("Avg Word Count", f"{int(avg_words):,}")
    
    with col3:
        avg_confidence = sum(m.confidence_score for m in metadata_list) / len(metadata_list)
        st.metric("Avg Confidence", f"{avg_confidence:.2f}")
    
    with col4:
        total_reading_time = sum(m.reading_time_minutes for m in metadata_list)
        st.metric("Total Reading Time", f"{total_reading_time} min")
    
    # Charts
    col1, col2 = st.columns(2)
    
    with col1:
        # Document types pie chart
        doc_types = [m.document_type for m in metadata_list]
        type_counts = pd.Series(doc_types).value_counts()
        
        fig = px.pie(values=type_counts.values, names=type_counts.index, 
                    title="Document Types Distribution")
        st.plotly_chart(fig, use_container_width=True)
    
    with col2:
        # Sentiment distribution
        sentiments = [m.sentiment for m in metadata_list]
        sentiment_counts = pd.Series(sentiments).value_counts()
        
        fig = px.bar(x=sentiment_counts.index, y=sentiment_counts.values,
                    title="Sentiment Distribution", 
                    color=sentiment_counts.values,
                    color_continuous_scale="RdYlBu")
        st.plotly_chart(fig, use_container_width=True)

def display_metadata_card(metadata, index):
    """Display a metadata card with key information"""
    with st.expander(f"📄 {metadata.filename} - {metadata.title}", expanded=False):
        col1, col2 = st.columns([2, 1])
        
        with col1:
            st.write("**Summary:**", metadata.summary)
            st.write("**Topics:**", ", ".join(metadata.primary_topics))
            st.write("**Keywords:**", ", ".join(metadata.keywords[:5]))
            
            if hasattr(metadata, 'entities') and metadata.entities:
                st.write("**Entities:**")
                for entity_type, entities in metadata.entities.items():
                    if entities:
                        st.write(f"  - {entity_type}: {', '.join(entities[:3])}")
        
        with col2:
            st.metric("Word Count", f"{metadata.word_count:,}")
            st.metric("Confidence", f"{metadata.confidence_score:.2f}")
            st.metric("Reading Time", f"{metadata.reading_time_minutes} min")
            st.metric("Document Type", metadata.document_type)

def main():
    """Main Streamlit application"""
    
    # Page configuration
    st.set_page_config(
        page_title="Automated Metadata Generator",
        page_icon="📄",
        layout="wide",
        initial_sidebar_state="expanded"
    )
    
    # Custom CSS for better styling
    st.markdown("""
    <style>
    .main-header {
        text-align: center;
        color: #2E86AB;
        padding: 1rem 0;
        border-bottom: 2px solid #f0f2f6;
        margin-bottom: 2rem;
    }
    
    .metric-card {
        background: linear-gradient(90deg, #667eea 0%, #764ba2 100%);
        padding: 1rem;
        border-radius: 10px;
        color: white;
        text-align: center;
        margin: 0.5rem 0;
    }
    
    .info-box {
        background: #f8f9fa;
        padding: 1rem;
        border-radius: 8px;
        border-left: 4px solid #2E86AB;
        margin: 1rem 0;
    }
    
    .success-box {
        background: #d4edda;
        color: #155724;
        padding: 1rem;
        border-radius: 8px;
        border: 1px solid #c3e6cb;
        margin: 1rem 0;
    }
    </style>
    """, unsafe_allow_html=True)
    
    # Header
    st.markdown("""
    <div class="main-header">
        <h1>🤖 Automated Metadata Generation System</h1>
        <p>Upload documents and get intelligent metadata automatically generated using advanced NLP analysis</p>
    </div>
    """, unsafe_allow_html=True)
    
    # Initialize components
    processor, analyzer, generator = initialize_components()
    
    # Initialize session state
    if 'processed_documents' not in st.session_state:
        st.session_state.processed_documents = []
        
    # Sidebar
    with st.sidebar:
       
        
        st.markdown("## 🚀 Features")
        st.markdown("""
        - **Multi-format Support**: PDF, DOCX, TXT
        - **OCR Capability**: Scanned documents
        - **Smart Analysis**: Topics, entities, sentiment
        - **Structured Output**: JSON, XML, YAML
        - **Batch Processing**: Multiple files
        """)
        
        st.markdown("### 📊 Analytics")
        if st.session_state.processed_documents:
            total_docs = len(st.session_state.processed_documents)
            total_words = sum(doc.word_count for doc in st.session_state.processed_documents)
            st.metric("Processed Documents", total_docs)
            st.metric("Total Words Analyzed", f"{total_words:,}")
        else:
            st.info("Upload documents to see analytics")
        
        # Settings
        st.markdown("### ⚙️ Settings")
        export_format = st.selectbox("Export Format", ["JSON", "XML", "YAML"])
        include_confidence = st.checkbox("Include Confidence Scores", value=True)
        batch_mode = st.checkbox("Batch Processing Mode", value=False)
    
    # Main content area
    tab1, tab2, tab3, tab4 = st.tabs(["📤 Upload & Process", "📊 Analytics", "📋 Results", "⚙️ Settings"])
    
    with tab1:
        st.header("Document Upload and Processing")
        
        # File upload section
        st.subheader("1. Upload Documents")
        uploaded_files = st.file_uploader(
            "Choose files to process",
            type=['pdf', 'docx', 'txt', 'doc'],
            accept_multiple_files=batch_mode,
            help="Supported formats: PDF, DOCX, TXT, DOC"
        )
        
        # Custom tags input
        st.subheader("2. Custom Configuration")
        col1, col2 = st.columns(2)
        
        with col1:
            custom_tags = st.text_input(
                "Custom Tags (comma-separated)",
                placeholder="e.g., important, quarterly-report, 2024"
            )
        
        with col2:
            processing_options = st.multiselect(
                "Processing Options",
                ["Extract Images", "Deep Topic Analysis", "Entity Linking", "OCR Processing"],
                default=["Deep Topic Analysis"]
            )
        
        # Process button
        if st.button("🚀 Process Documents", type="primary", use_container_width=True):
            if uploaded_files:
                files_to_process = uploaded_files if isinstance(uploaded_files, list) else [uploaded_files]
                
                progress_bar = st.progress(0)
                status_text = st.empty()
                
                for i, uploaded_file in enumerate(files_to_process):
                    status_text.text(f"Processing {uploaded_file.name}...")
                    
                    try:
                        # Save uploaded file temporarily
                        temp_dir = Path("temp_uploads")
                        temp_dir.mkdir(exist_ok=True)
                        
                        temp_file_path = temp_dir / uploaded_file.name
                        with open(temp_file_path, "wb") as f:
                            f.write(uploaded_file.getbuffer())
                        
                        # Process the saved file
                        doc_result = processor.process_document(str(temp_file_path))
                        nlp_result = analyzer.analyze_document(doc_result['processed_text'], uploaded_file.name)
                        
                        # Clean up temporary file
                        temp_file_path.unlink()
                        
                    except Exception as e:
                        st.error(f"Error processing {uploaded_file.name}: {str(e)}")
                        continue
                    
                    # Generate metadata
                    tags_list = [tag.strip() for tag in custom_tags.split(',')] if custom_tags else None
                    metadata = generator.generate_metadata(doc_result, nlp_result, tags_list)
                    
                    # Store in session state
                    st.session_state.processed_documents.append(metadata)
                    
                    # Update progress
                    progress_bar.progress((i + 1) / len(files_to_process))
                
                status_text.text("✅ Processing complete!")
                st.success(f"Successfully processed {len(files_to_process)} document(s)!")
                
                # Auto-switch to results tab
                st.rerun()
            else:
                st.warning("Please upload at least one document to process.")
    
    with tab2:
        st.header("Analytics Dashboard")
        if st.session_state.processed_documents:
            create_analytics_dashboard(st.session_state.processed_documents)
        else:
            st.info("No documents processed yet. Upload and process documents to see analytics.")
    
    with tab3:
        st.header("Processing Results")
        
        if st.session_state.processed_documents:
            # Display options
            col1, col2, col3 = st.columns(3)
            with col1:
                view_mode = st.selectbox("View Mode", ["Cards", "Table", "JSON"])
            with col2:
                sort_by = st.selectbox("Sort By", ["Filename", "Confidence", "Word Count", "Processing Date"])
            with col3:
                if st.button("🗑️ Clear All Results"):
                    st.session_state.processed_documents = []
                    st.rerun()
            
            # Display results based on view mode
            if view_mode == "Cards":
                for i, metadata in enumerate(st.session_state.processed_documents):
                    display_metadata_card(metadata, i)
            
            elif view_mode == "Table":
                # Create DataFrame for table view
                table_data = []
                for metadata in st.session_state.processed_documents:
                    table_data.append({
                        'Filename': metadata.filename,
                        'Title': metadata.title,
                        'Document Type': metadata.document_type,
                        'Word Count': metadata.word_count,
                        'Confidence': metadata.confidence_score,
                        'Sentiment': metadata.sentiment,
                        'Topics': ', '.join(metadata.primary_topics[:3])
                    })
                
                df = pd.DataFrame(table_data)
                st.dataframe(df, use_container_width=True)
            
            elif view_mode == "JSON":
                for i, metadata in enumerate(st.session_state.processed_documents):
                    with st.expander(f"JSON - {metadata.filename}"):
                        json_output = generator.export_metadata(metadata, 'json')
                        st.code(json_output, language='json')
            
            # Bulk export
            st.subheader("📥 Export Results")
            col1, col2 = st.columns(2)
            
            with col1:
                if st.button("📄 Export All as JSON"):
                    all_metadata = []
                    for metadata in st.session_state.processed_documents:
                        json_data = json.loads(generator.export_metadata(metadata, 'json'))
                        all_metadata.append(json_data)
                    
                    output = json.dumps(all_metadata, indent=2)
                    st.download_button(
                        "Download JSON",
                        output,
                        file_name="metadata_export.json",
                        mime="application/json"
                    )
            
            with col2:
                if st.button("📊 Export as CSV"):
                    table_data = []
                    for metadata in st.session_state.processed_documents:
                        table_data.append({
                            'filename': metadata.filename,
                            'title': metadata.title,
                            'summary': metadata.summary,
                            'document_type': metadata.document_type,
                            'word_count': metadata.word_count,
                            'confidence_score': metadata.confidence_score,
                            'sentiment': metadata.sentiment,
                            'topics': '|'.join(metadata.primary_topics),
                            'keywords': '|'.join(metadata.keywords)
                        })
                    
                    df = pd.DataFrame(table_data)
                    csv = df.to_csv(index=False)
                    st.download_button(
                        "Download CSV",
                        csv,
                        file_name="metadata_export.csv",
                        mime="text/csv"
                    )
        
        else:
            st.info("No results to display. Process some documents first!")
    
    with tab4:
        st.header("Advanced Settings")
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.subheader("🔧 Processing Configuration")
            
            confidence_threshold = st.slider(
                "Confidence Threshold",
                min_value=0.0,
                max_value=1.0,
                value=0.5,
                step=0.1,
                help="Minimum confidence score for topic classification"
            )
            
            max_keywords = st.number_input(
                "Maximum Keywords",
                min_value=5,
                max_value=50,
                value=20,
                help="Maximum number of keywords to extract"
            )
            
            enable_ocr = st.checkbox("Enable OCR for Scanned Documents", value=True)
            enable_entity_linking = st.checkbox("Enable Entity Linking", value=False)
        
        with col2:
            st.subheader("📊 Output Configuration")
            
            include_metadata = st.multiselect(
                "Include in Output",
                ["File Statistics", "Language Features", "Document Structure", "Processing Logs"],
                default=["File Statistics", "Document Structure"]
            )
            
            output_language = st.selectbox(
                "Summary Language",
                ["Auto-detect", "English", "Spanish", "French", "German"],
                index=0
            )
            
            st.subheader("💾 Data Management")
            if st.button("🧹 Clear Processing Cache"):
                st.cache_resource.clear()
                st.success("Cache cleared successfully!")
    
    # Footer
    st.markdown("---")
    st.markdown("""
    <div style='text-align: center; color: #666; padding: 1rem;'>
        <p>🤖 Automated Metadata Generation System v1.0</p>
        <p>Built with Streamlit • Powered by Advanced NLP</p>
    </div>
    """, unsafe_allow_html=True)

if __name__ == "__main__":
    main()
